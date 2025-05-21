import requests
import json
import urllib3
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
from concurrent.futures import ThreadPoolExecutor, as_completed

# Disable SSL warnings
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# Define API endpoints
base_url = "https://app.plantagents.org"
plant_endpoints = {
    "find_by_name": f"{base_url}/Plant/FindByName",
    "find_vendors_by_plant_name": f"{base_url}/Plant/FindVendorsForPlantName"
}

# API key
api_key = "d7b141d6-c12e-4fb1-bedd-8868a0a6f4e6"

# Set headers
headers = {
    "Authorization": f"Bearer {api_key}",
    "Content-Type": "application/json",
    "Accept": "application/json"
}

def fetch_plants(plant_name):
    try:
        print(f"\nSearching for plants with name: {plant_name}")
        params = {"plantName": plant_name}
        response = requests.get(plant_endpoints["find_by_name"], headers=headers, params=params, verify=False)
        print(f"Response status: {response.status_code}")
        
        if response.status_code == 200:
            data = response.json()
            if isinstance(data, list):
                return data
            return []
        else:
            print(f"Error searching plants: {response.text[:200]}")
            return None
    except Exception as e:
        print(f"Error: {e}")
        return None

def fetch_vendors_for_plant_and_zip(plant_name, zip_code):
    try:
        params = {
            "plantName": plant_name,
            "zipCode": zip_code,
            "radius": 25,  # Reduced radius to 25 miles
            "limit": 100
        }
        print(f"Searching for {plant_name} near {zip_code} (radius: 25 miles)")
        response = requests.get(plant_endpoints["find_vendors_by_plant_name"], headers=headers, params=params, verify=False)
        
        if response.status_code == 200:
            vendors = response.json()
            if vendors:
                print(f"Found {len(vendors)} vendors for {plant_name} near {zip_code}")
            return vendors or []
        else:
            print(f"Error response for {plant_name} near {zip_code}: {response.status_code}")
            return []
    except Exception as e:
        print(f"Error searching vendors for {plant_name} near {zip_code}: {e}")
        return []

def fetch_vendors_for_plant(plant_name, zip_codes=None):
    if zip_codes is None:
        # Philadelphia and surrounding area zip codes
        zip_codes = [
            "19103",  # Center City
            "19102",  # Center City West
            "19104",  # University City
            "19106",  # Old City
            "19107",  # Washington Square West
            "19123",  # Northern Liberties
            "19125",  # Fishtown
            "19130",  # Fairmount
            "19147",  # Queen Village
            "19148",  # South Philadelphia
            "19129",  # East Falls
            "19128",  # Roxborough
            "19119",  # Mount Airy
            "19118",  # Chestnut Hill
            "19144"   # Germantown
        ]
    
    all_vendors = []
    seen_vendors = set()
    
    # Use ThreadPoolExecutor to search multiple zip codes simultaneously
    with ThreadPoolExecutor(max_workers=5) as executor:
        # Create a future for each zip code
        future_to_zip = {
            executor.submit(fetch_vendors_for_plant_and_zip, plant_name, zip_code): zip_code 
            for zip_code in zip_codes
        }
        
        # Process results as they complete
        for future in as_completed(future_to_zip):
            zip_code = future_to_zip[future]
            try:
                vendors = future.result()
                if vendors:
                    for vendor in vendors:
                        vendor_id = vendor.get('storeName', '') + vendor.get('address', '')
                        if vendor_id not in seen_vendors:
                            seen_vendors.add(vendor_id)
                            all_vendors.append(vendor)
            except Exception as e:
                print(f"Error processing results for {plant_name} near {zip_code}: {e}")
    
    print(f"Found {len(all_vendors)} unique vendors for {plant_name}")
    return all_vendors

def fetch_all_plants():
    """Fetch all plants from the API using multiple search terms"""
    print("Fetching all plants from the API...")
    all_plants = set()
    
    # Search terms that should return native plants
    search_terms = [
        "flower",
        "plant",
        "tree",
        "shrub",
        "grass",
        "vine",
        "fern"
    ]
    
    for term in search_terms:
        print(f"\nSearching with term: {term}")
        try:
            params = {"plantName": term}
            response = requests.get(
                plant_endpoints["find_by_name"],
                params=params,
                headers=headers,
                verify=False
            )
            print(f"Response status: {response.status_code}")
            
            if response.status_code == 200:
                plants = response.json()
                if isinstance(plants, list):
                    for plant in plants:
                        if isinstance(plant, dict):
                            scientific_name = plant.get('scientificName', '')
                            common_name = plant.get('commonName', '')
                            if not scientific_name and not common_name:
                                # Try to extract from blurb if available
                                blurb = plant.get('blurb', '')
                                if blurb:
                                    # Extract first sentence as common name
                                    common_name = blurb.split('.')[0].strip()
                            if scientific_name or common_name:
                                all_plants.add((scientific_name, common_name))
                    print(f"Found {len(plants)} plants for term '{term}'")
                else:
                    print(f"Unexpected response format for term '{term}': {response.text[:200]}")
            else:
                print(f"Error response for term '{term}': {response.text[:200]}")
        except Exception as e:
            print(f"Error searching for term '{term}': {str(e)}")
    
    print(f"\nFound {len(all_plants)} unique plants\n")
    return list(all_plants)

def create_plant_document(plants_data):
    doc = Document()
    
    # Add title
    title = doc.add_heading('Native Plants and Nurseries in Philadelphia', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add introduction
    intro = doc.add_paragraph()
    intro.add_run("This document lists native plants and their availability at nurseries across Philadelphia.")
    
    # Process each plant
    print("\nGathering vendor information for all plants...")
    plants_with_vendors = []
    total_plants = len(plants_data)
    
    for index, plant in enumerate(plants_data, 1):
        print(f"\nProcessing plant {index}/{total_plants}")
        scientific_name, common_name = plant
        
        # First try searching with common name
        vendors = []
        if common_name:
            vendors = fetch_vendors_for_plant(common_name)
        
        # If no vendors found with common name, try scientific name
        if not vendors and scientific_name:
            vendors = fetch_vendors_for_plant(scientific_name)
        
        plants_with_vendors.append({
            'scientific_name': scientific_name,
            'common_name': common_name,
            'vendors': vendors,
            'vendor_count': len(vendors)
        })
    
    # Sort plants by number of vendors (most available first)
    plants_with_vendors.sort(key=lambda x: x['vendor_count'], reverse=True)
    
    # Add plants to document
    for plant_data in plants_with_vendors:
        # Add plant heading
        heading = doc.add_heading(level=1)
        heading.add_run(f"{plant_data['common_name']} ({plant_data['scientific_name']})")
        
        # Add vendor information
        if plant_data['vendors']:
            doc.add_paragraph("Available at the following nurseries:")
            for vendor in plant_data['vendors']:
                vendor_info = doc.add_paragraph(style='List Bullet')
                vendor_info.add_run(f"{vendor.get('storeName', 'Unknown Store')} - {vendor.get('address', 'No address available')}")
        else:
            doc.add_paragraph("No nurseries found in Philadelphia area.")
        
        # Add spacing between plants
        doc.add_paragraph()
    
    # Save the document
    doc.save('philadelphia_plant_nursery_list.docx')
    print("\nWord document created successfully: philadelphia_plant_nursery_list.docx")

def main():
    # Fetch all plants
    all_plants = fetch_all_plants()
    
    if all_plants:
        print(f"\nProcessing {len(all_plants)} plants")
        create_plant_document(all_plants)
    else:
        print("\nNo plants found. Please check the API connection.")

if __name__ == "__main__":
    main() 